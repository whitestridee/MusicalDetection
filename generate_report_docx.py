from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "Отчет_поиск_похожих_музыкальных_композиций_v3.docx"
FIGURES_DIR = ROOT / "report_figures"


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT, size: int = 12) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "B7B7B7")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def load_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/timesbd.ttf"),
                Path("C:/Windows/Fonts/arialbd.ttf"),
                Path("C:/Windows/Fonts/calibrib.ttf"),
            ]
        )
    else:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/times.ttf"),
                Path("C:/Windows/Fonts/arial.ttf"),
                Path("C:/Windows/Fonts/calibri.ttf"),
            ]
        )
    for candidate in candidates:
        if candidate.exists():
            try:
                return ImageFont.truetype(str(candidate), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def draw_centered_multiline(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill, spacing: int = 10) -> None:
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = box[0] + ((box[2] - box[0]) - text_width) / 2
    y = box[1] + ((box[3] - box[1]) - text_height) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=spacing, align="center")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_number_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def add_experiment_table(doc: Document, title: str, rows: list[tuple[str, str, str]]) -> None:
    add_body_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].width = Cm(7)
    header[1].width = Cm(4)
    header[2].width = Cm(4)
    set_cell_text(header[0], "Трек", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[1], "Similarity", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[2], "Confidence", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    for name, similarity, confidence in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(7)
        cells[1].width = Cm(4)
        cells[2].width = Cm(4)
        set_cell_text(cells[0], name, size=12)
        set_cell_text(cells[1], similarity, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], confidence, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_module_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].width = Cm(4.2)
    header[1].width = Cm(4.8)
    header[2].width = Cm(7.0)
    set_cell_text(header[0], "Модуль", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[1], "Файл", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[2], "Назначение", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Загрузка сигнала", "audio.py", "Чтение WAV-файлов, перевод в mono и ресемплинг."),
        ("Признаки", "features.py", "Расчет спектральных признаков, mel-энергий, MFCC и delta-MFCC."),
        ("Индекс", "index.py", "Построение матрицы embeddings, стандартизация и сохранение статистик."),
        ("Поиск", "search.py", "Косинусная близость, ранжирование кандидатов и расчет confidence."),
        ("CLI-индексация", "build_index.py", "Оффлайн-подготовка базы и сохранение индекса на диск."),
        ("HTTP API", "server.py", "Точка входа веб-сервиса и запуск сервера."),
        ("MVC", "mvc_app/*", "Модель поиска, контроллеры маршрутов и HTML-представления."),
    ]
    for row in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(4.8)
        cells[2].width = Cm(7.0)
        set_cell_text(cells[0], row[0], size=12)
        set_cell_text(cells[1], row[1], size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[2], row[2], size=12)


def add_dataset_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].width = Cm(8)
    header[1].width = Cm(8)
    set_cell_text(header[0], "Жанр", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header[1], "Количество файлов", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for genre in ["blues", "classical", "country", "disco", "hiphop", "jazz", "metal", "pop", "reggae", "rock"]:
        cells = table.add_row().cells
        cells[0].width = Cm(8)
        cells[1].width = Cm(8)
        set_cell_text(cells[0], genre, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], "10", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)


def make_figures() -> tuple[Path, Path]:
    FIGURES_DIR.mkdir(exist_ok=True)
    title_font = load_font(44, bold=True)
    body_font = load_font(34, bold=True)
    small_font = load_font(28, bold=False)

    pipeline = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(pipeline)
    accent = (72, 92, 153)
    fill = (232, 238, 252)
    text = (33, 33, 33)

    draw.text((70, 40), "Пайплайн поиска похожих музыкальных композиций", font=title_font, fill=text)

    boxes = [
        ((70, 300, 360, 470), "Аудио\nзапрос"),
        ((430, 300, 760, 470), "Предобработка\nмоно, 16 кГц,\nнормализация"),
        ((830, 300, 1180, 470), "Извлечение\nпризнаков\nspectral + MFCC"),
        ((1250, 300, 1610, 470), "Стандартизация\nи нормализация\nвектора"),
        ((430, 640, 760, 810), "База треков\nGTZAN\n100 WAV"),
        ((830, 640, 1180, 810), "Индекс\nэмбеддингов\nNumPy"),
        ((1250, 640, 1610, 810), "Top-N поиск\nпо косинусной\nблизости"),
    ]

    for box, label in boxes:
        draw.rounded_rectangle(box, radius=24, fill=fill, outline=accent, width=5)
        draw_centered_multiline(draw, box, label, body_font, text, spacing=12)

    arrows = [
        ((360, 385), (430, 385)),
        ((760, 385), (830, 385)),
        ((1180, 385), (1250, 385)),
        ((595, 640), (595, 470)),
        ((1005, 640), (1005, 470)),
        ((1430, 640), (1430, 470)),
        ((760, 725), (830, 725)),
        ((1180, 725), (1250, 725)),
    ]
    for start, end in arrows:
        draw.line([start, end], fill=accent, width=7)
        draw.polygon(
            [(end[0], end[1]), (end[0] - 22, end[1] - 12), (end[0] - 22, end[1] + 12)],
            fill=accent,
        )

    draw.text((70, 900), "Оффлайн-индексация и онлайн-поиск используют одинаковое извлечение признаков.", font=small_font, fill=(70, 70, 70))

    pipeline_path = FIGURES_DIR / "pipeline.png"
    pipeline.save(pipeline_path)

    mvc = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(mvc)
    accent2 = (145, 74, 64)
    fill2 = (249, 234, 230)
    green = (88, 123, 96)
    green_fill = (231, 243, 233)

    draw.text((70, 40), "MVC-структура веб-приложения", font=title_font, fill=text)

    view_box = (90, 180, 520, 500)
    controller_box = (685, 180, 1115, 500)
    model_box = (1280, 180, 1710, 500)
    core_box = (685, 640, 1115, 900)

    draw.rounded_rectangle(view_box, radius=26, fill=fill2, outline=accent2, width=5)
    draw_centered_multiline(draw, view_box, "View\nHTML-страница\nформа загрузки\nтаблица результатов", body_font, text, spacing=12)

    draw.rounded_rectangle(controller_box, radius=26, fill=(242, 239, 250), outline=(95, 78, 145), width=5)
    draw_centered_multiline(draw, controller_box, "Controller\nмаршруты GET/POST\nпарсинг multipart\nвалидация запроса", body_font, text, spacing=12)

    draw.rounded_rectangle(model_box, radius=26, fill=green_fill, outline=green, width=5)
    draw_centered_multiline(draw, model_box, "Model\nпоиск похожих\nтреков\nработа с индексом", body_font, text, spacing=12)

    draw.rounded_rectangle(core_box, radius=26, fill=(247, 247, 247), outline=(110, 110, 110), width=5)
    draw_centered_multiline(draw, core_box, "audio_similarity\nfeatures.py\nindex.py\nsearch.py", body_font, text, spacing=14)

    links = [
        ((520, 340), (685, 340)),
        ((1115, 340), (1280, 340)),
        ((900, 500), (900, 640)),
    ]
    for start, end in links:
        draw.line([start, end], fill=(80, 80, 80), width=7)
        draw.polygon(
            [(end[0], end[1]), (end[0] - 14, end[1] - 22), (end[0] + 14, end[1] - 22)],
            fill=(80, 80, 80),
        )

    draw.text((70, 960), "Разделение на Model, View и Controller упрощает поддержку сервиса и развитие интерфейса.", font=small_font, fill=(70, 70, 70))

    mvc_path = FIGURES_DIR / "mvc.png"
    mvc.save(mvc_path)
    return pipeline_path, mvc_path


def build_report() -> Path:
    pipeline_path, mvc_path = make_figures()

    doc = Document()
    configure_document(doc)

    org_table = doc.add_table(rows=1, cols=2)
    org_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    org_table.autofit = False
    org_table.cell(0, 0).width = Cm(2.2)
    org_table.cell(0, 1).width = Cm(13.8)
    org_table.cell(0, 0).text = ""
    set_cell_text(
        org_table.cell(0, 1),
        "Министерство науки и высшего образования Российской Федерации\n"
        "Федеральное государственное бюджетное образовательное учреждение\n"
        "высшего образования\n"
        "«Московский государственный технический университет\n"
        "имени Н.Э. Баумана\n"
        "(национальный исследовательский университет)»\n"
        "(МГТУ им. Н.Э. Баумана)",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
    )

    for text in [
        "",
        "ФАКУЛЬТЕТ Информатика и системы управления",
        "",
        "КАФЕДРА Технологии искусственного интеллекта (ИУ12)",
        "",
        "НАПРАВЛЕНИЕ ПОДГОТОВКИ 09.04.01 Информатика и вычислительная техника",
        "",
        "",
        "Отчет",
        "",
        "Разработка системы поиска похожих музыкальных композиций",
        "",
        "Дисциплина: Цифровая обработка сигналов",
        "",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if text == "Отчет":
            run.font.size = Pt(24)
            run.bold = True
        elif text == "Разработка системы поиска похожих музыкальных композиций":
            run.font.size = Pt(20)
            run.bold = True
        else:
            run.font.size = Pt(14)

    sign_table = doc.add_table(rows=5, cols=5)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    widths = [Cm(3.2), Cm(3.0), Cm(2.7), Cm(3.2), Cm(4.0)]
    for row in sign_table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    sign_rows = [
        ["Студент", "ИУ12-31М", "", "", "Б.Б. Дедич"],
        ["", "(Группа)", "", "(Подпись, дата)", "(И.О. Фамилия)"],
        ["", "", "", "", ""],
        ["Преподаватель", "", "", "", "Д. Ю. Евсюков"],
        ["", "", "", "(Подпись, дата)", "(И.О. Фамилия)"],
    ]
    for row_idx, values in enumerate(sign_rows):
        for col_idx, value in enumerate(values):
            set_cell_text(
                sign_table.cell(row_idx, col_idx),
                value,
                size=11 if row_idx in (1, 4) else 12,
                align=WD_ALIGN_PARAGRAPH.CENTER if col_idx != 4 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_table_borders(sign_table)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nМосква, 2026")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)

    doc.add_page_break()

    add_heading(doc, "1. Цель работы")
    add_body_paragraph(doc, "Целью работы является разработка программного сервиса на языке Python, который принимает на вход музыкальный трек или его фрагмент, извлекает из аудиосигнала устойчивое векторное представление и возвращает список наиболее похожих композиций из заранее подготовленной базы данных. В рамках работы необходимо не только реализовать алгоритм поиска ближайших соседей, но и построить полный прикладной пайплайн: от подготовки данных и индексации базы до пользовательского интерфейса и анализа качества результатов.")
    add_body_paragraph(doc, "Практическая значимость работы заключается в том, что задачи поиска по аудио широко применяются в музыкальных каталогах, рекомендательных системах, сервисах распознавания контента, мультимедийных архивах и системах интеллектуального анализа медиа. При этом для учебного проекта особенно важен баланс между качеством результата, интерпретируемостью метода и сложностью реализации. Поэтому в отчете рассматривается baseline-подход, который можно воспроизвести на стандартном наборе инструментов Python и затем постепенно усиливать более сложными моделями.")

    add_heading(doc, "2. Постановка задачи")
    add_body_paragraph(doc, "В соответствии с заданием необходимо построить сервис Audio Similarity Search, работающий с базой порядка 100-1000 треков. В качестве исходной коллекции в проекте использован датасет GTZAN Genre Collection, из которого была сформирована рабочая подвыборка из 100 WAV-файлов: по 10 треков для каждого жанра blues, classical, country, disco, hiphop, jazz, metal, pop, reggae и rock. Такой выбор позволил собрать сбалансированную демонстрационную базу и обеспечить достаточно быстрый пересчет индекса при многочисленных экспериментах.")
    add_body_paragraph(doc, "Формально решаемая задача сводится к поиску элементов x_i из базы D = {x_1, x_2, ..., x_N}, которые максимизируют меру сходства s(q, x_i) между запросом q и векторными представлениями треков. Если обозначить через f(.) функцию извлечения признаков, то для каждого трека вычисляется embedding z_i = f(x_i), а для запроса строится вектор z_q = f(q). Далее по множеству {z_i} выполняется поиск ближайших соседей по косинусной близости.")
    add_formula(doc, "sim(z_q, z_i) = (z_q · z_i) / (||z_q|| ||z_i||)")
    add_body_paragraph(doc, "Важное требование к системе состоит в том, чтобы к запросному файлу применялась точно такая же предобработка, что и к базе. Это означает одинаковую частоту дискретизации, перевод в моно, нормализацию, разбиение на окна, извлечение спектральных признаков и последующую стандартизацию. Нарушение этого требования делает сравнение векторов некорректным и резко снижает устойчивость поиска.")

    add_heading(doc, "3. Краткие теоретические сведения")
    add_heading(doc, "3.1. Поиск похожих аудиосигналов", level=2)
    add_body_paragraph(doc, "Поиск похожих музыкальных композиций относится к классу задач content-based retrieval, где решение принимается не на основе текстовых метаданных, а по свойствам самого сигнала. В простейшем случае система анализирует временную структуру и спектральное распределение энергии, а затем представляет трек фиксированным вектором признаков. Чем ближе два вектора в выбранном пространстве, тем выше вероятность, что аудио обладает сходным тембром, ритмической организацией, частотной структурой или общей акустической картиной.")
    add_body_paragraph(doc, "Для музыки особенно важны признаки, связанные с коротковременным спектром. Аудиосигнал является нестационарным, поэтому его анализ обычно выполняется по небольшим перекрывающимся окнам. В каждом окне можно оценить энергию, частоту пересечений нуля, спектральный центроид, ширину спектра, rolloff, flatness и другие величины. Далее статистики усредняются, образуя вектор фиксированной длины. Такой подход удобен как baseline, поскольку он интерпретируем, недорог в вычислении и не требует обучения большой нейросети.")
    add_heading(doc, "3.2. Мел-шкала и MFCC", level=2)
    add_body_paragraph(doc, "Чтобы приблизить модель к особенностям человеческого слуха, в задачах музыкального и речевого анализа широко используется мел-шкала. Она нелинейно перераспределяет ось частот: низкие частоты описываются подробнее, а высокие — грубее. После применения набора треугольных мел-фильтров к спектру сигнала получается компактное описание распределения энергии по полосам. Логарифмирование этих энергий подчеркивает относительные изменения, а дискретное косинусное преобразование дает коэффициенты MFCC, которые давно считаются стандартом в аудиоанализе.")
    add_body_paragraph(doc, "В данной работе к исходным спектральным признакам были добавлены лог-мел-энергии, средние и стандартные отклонения MFCC, а также первые разности (delta-MFCC). Такое решение оказалось полезным по двум причинам. Во-первых, embedding стал лучше различать записи по тембровой структуре. Во-вторых, исчезла слишком сильная концентрация косинусных расстояний вблизи единицы, которая наблюдалась у более простого baseline.")
    add_heading(doc, "3.3. Индексирование и калибровка сходства", level=2)
    add_body_paragraph(doc, "После вычисления embedding для каждого трека база превращается в матрицу признаков. Даже если число объектов пока невелико, полезно строить явный индекс и сохранять его на диск, чтобы не пересчитывать векторы заново при каждом запуске сервера. В нашем проекте использован точный поиск по матрице NumPy, поскольку при 100-1000 объектах он остается быстрым, прозрачен для отладки и не требует внешних зависимостей. Тем не менее архитектура решения позволяет позже заменить этот этап на FAISS или HNSW без изменения интерфейса.")
    add_body_paragraph(doc, "Дополнительная проблема baseline-подходов состоит в плохой калибровке similarity. Если признаки имеют положительное смещение и похожий масштаб, то косинусная близость для большинства пар треков оказывается аномально высокой. Чтобы устранить этот эффект, в работе применяется стандартизация признаков по всей базе: из каждого измерения вычитается среднее значение и выполняется деление на стандартное отклонение. После этого и база, и запрос нормализуются по L2-норме, а confidence вычисляется не линейно, а как функция того, насколько сильно результат выделяется на фоне распределения остальных кандидатов.")

    add_heading(doc, "4. Описание данных и подготовка датасета")
    add_body_paragraph(doc, "Для демонстрации работы сервиса использован открытый набор GTZAN Genre Collection, содержащий 1000 аудиофрагментов по 30 секунд в десяти жанровых категориях. Исходный датасет хорошо подходит для учебных проектов, потому что уже организован по каталогам, имеет единый формат хранения и часто используется как ориентир в задачах музыкальной классификации. Из полного набора была подготовлена компактная база из 100 файлов, достаточная для тестирования индексации, сервиса и пользовательского интерфейса.")
    add_body_paragraph(doc, "На этапе подготовки базы были выполнены следующие действия:")
    add_bullet_list(doc, [
        "скачивание полного архива GTZAN и распаковка в рабочую директорию проекта;",
        "формирование сбалансированной подвыборки: по 10 WAV-файлов на жанр;",
        "сохранение файлов в каталоге dataset с жанровой структурой подпапок;",
        "создание вспомогательного индекса metadata.json и матрицы embeddings.npy для ускоренного доступа."
    ])
    add_body_paragraph(doc, "Состав используемой подвыборки приведен в таблице 1.")
    add_dataset_table(doc)
    add_body_paragraph(doc, "Выбранная подвыборка позволяет проверить как точность поиска по самому треку, так и устойчивость сервиса к внешним данным, например к записи с диктофона. Важный вывод этапа подготовки состоит в том, что качество retrieval сильно зависит не только от алгоритма, но и от домена базы: студийные музыкальные фрагменты и бытовая запись с микрофона относятся к разным акустическим условиям, поэтому даже хороший алгоритм не обязан возвращать интуитивно ожидаемый жанр.")

    add_heading(doc, "5. Архитектура приложения")
    add_body_paragraph(doc, "Программная система была реализована как набор модулей Python, разделенных по ответственности. Базовый пакет audio_similarity отвечает за работу с сигналом, извлечение признаков, формирование индекса и выполнение поиска. Веб-сервис построен поверх простой HTTP-архитектуры и дополнительно организован в стиле MVC, где модель инкапсулирует алгоритм поиска, контроллеры обрабатывают маршруты и загрузку файлов, а представления формируют HTML-страницы и таблицы результатов.")
    add_body_paragraph(doc, "На рисунке 1 показан общий технологический конвейер приложения. Он объединяет оффлайн-этап индексации базы и онлайн-этап обработки запроса. Такой подход удобен тем, что тяжелые операции по подготовке коллекции выполняются один раз, а пользователь при запросе взаимодействует только с готовым индексом.")
    doc.add_picture(str(pipeline_path), width=Cm(15.5))
    add_caption(doc, "Рисунок 1 — Общий пайплайн сервиса поиска похожих музыкальных композиций")
    add_body_paragraph(doc, "С точки зрения пользовательского сценария работа приложения выглядит следующим образом. Сначала оператор подготавливает базу WAV-файлов и запускает скрипт build_index.py, который строит индекс embeddings.npy и сохраняет metadata.json. Затем запускается сервер server.py. Пользователь открывает веб-интерфейс, загружает WAV-файл и указывает число результатов top-k. Контроллер временно сохраняет файл, передает его модели поиска, а затем в браузере отображается отсортированная таблица кандидатов с similarity и confidence.")
    add_body_paragraph(doc, "Для удобства сопровождения было важно поддерживать прозрачную модульную структуру проекта. Основные компоненты программной системы перечислены в таблице 2.")
    add_module_table(doc)

    add_heading(doc, "6. Реализация извлечения признаков")
    add_body_paragraph(doc, "Модуль audio.py отвечает за загрузку PCM WAV-файлов, перевод стерео в моно и приведение сигнала к диапазону [-1, 1]. Далее выполняется линейная ресемплизация до 16 кГц. Такое значение выбрано как компромисс между скоростью обработки и сохранением информации в музыкальном сигнале. После этого в features.py применяется легкий pre-emphasis, который усиливает высокочастотные компоненты и делает спектральные различия более заметными.")
    add_body_paragraph(doc, "Следующий этап — разбиение на окна длиной 2048 отсчетов с шагом 512 и умножение на окно Хэннинга. Для каждого окна вычисляются традиционные спектральные признаки: RMS-энергия, zero-crossing rate, спектральный центроид, ширина спектра, rolloff и flatness. Параллельно формируется сжатый лог-спектр по 64 полосам, а также мел-энергии, MFCC и delta-MFCC. Затем по всем временным кадрам берутся средние значения и стандартные отклонения, из которых формируется итоговый embedding фиксированной длины.")
    add_body_paragraph(doc, "Идея агрегирования средних и стандартных отклонений принципиально важна. С одной стороны, она позволяет работать с файлами произвольной длительности и всегда получать вектор одинакового размера. С другой стороны, такой подход неизбежно теряет точную временную структуру и поэтому не способен полноценно распознавать композицию как последовательность музыкальных событий. Именно поэтому текущая реализация корректнее всего описывается как система coarse similarity, а не как промышленный аудиофингерпринтинг.")

    add_heading(doc, "7. Построение индекса и поиск")
    add_body_paragraph(doc, "После извлечения сырых векторов по всей базе строится матрица признаков размерности N x d, где N — число треков, а d — размер embedding. По этой матрице рассчитываются вектор средних feature_mean и вектор стандартных отклонений feature_std. Затем матрица стандартизуется, после чего каждая строка нормализуется по евклидовой норме. В таком виде embeddings.npy уже можно использовать для быстрого скалярного произведения с вектором запроса.")
    add_formula(doc, "z'_i = (z_i - μ) / σ,    z''_i = z'_i / ||z'_i||")
    add_body_paragraph(doc, "Для query-файла применяются те же сохраненные статистики стандартизации, что и для базы. Это важный момент: если стандартизовать запрос независимо, то геометрия пространства признаков исказится и similarity перестанет быть сопоставимой. После этого выполняется точный поиск ближайших соседей, результаты сортируются по убыванию similarity, а confidence вычисляется через сигмоиду от z-оценки относительно распределения всех оценок для данного запроса. Благодаря этому confidence начинает отражать не абсолютное значение косинусной близости, а выраженность кандидата на фоне всей базы.")
    add_body_paragraph(doc, "Именно этот шаг позволил избавиться от ситуации, когда почти все треки имели similarity 0.99 и выше. После введения стандартизации диапазон оценок стал заметно шире: для диктофонного файла значения распределялись приблизительно от -0.78 до 0.80, что намного лучше отражает реальную неопределенность поиска.")

    add_heading(doc, "8. Веб-интерфейс и MVC-организация")
    add_body_paragraph(doc, "По требованию проекта к базовой консольной версии был добавлен HTTP-сервис и пользовательский веб-интерфейс с загрузкой файла. Для упрощения поддержки код был приведен к архитектуре MVC. Модуль mvc_app.models содержит SearchModel, который инкапсулирует взаимодействие с индексом и функцией similarity search. Модуль mvc_app.controllers обрабатывает GET/POST-запросы, парсит multipart/form-data и организует маршрут search-upload. Наконец, mvc_app.views отвечает за формирование HTML-страницы с формой загрузки и таблицей результатов.")
    doc.add_picture(str(mvc_path), width=Cm(15.5))
    add_caption(doc, "Рисунок 2 — MVC-структура веб-приложения")
    add_body_paragraph(doc, "Такое разделение удобно и для сопровождения, и для защиты работы. Во-первых, становится ясно, где расположена предметная логика, а где — код интерфейса. Во-вторых, приложение легко расширять: например, можно заменить HTML-страницу на другой frontend, не переписывая модель поиска. В-третьих, архитектура отвечает хорошей инженерной практике и делает проект ближе к реальным прикладным сервисам, а не к одноразовому скрипту.")

    add_heading(doc, "9. Руководство по запуску и использованию сервиса")
    add_body_paragraph(doc, "Проект рассчитан на воспроизводимый запуск в обычной файловой структуре. На первом этапе пользователь размещает WAV-файлы в каталоге dataset, после чего строит индекс командой build_index.py. Сервис сохраняет embeddings.npy, metadata.json и статистики feature_stats.npz. На втором этапе запускается сервер server.py, который поднимает HTTP-приложение на локальном адресе 127.0.0.1:8000. На третьем этапе пользователь открывает главную страницу, загружает файл и получает ранжированный список похожих треков.")
    add_body_paragraph(doc, "Последовательность запуска системы следующая:")
    add_number_list(doc, [
        "Подготовить папку dataset и поместить в нее WAV-файлы базы.",
        "Выполнить индексацию командой build_index.py --dataset .\\dataset --output .\\index_data.",
        "Запустить сервер server.py --index .\\index_data --host 127.0.0.1 --port 8000.",
        "Открыть браузер по адресу http://127.0.0.1:8000/.",
        "Выбрать WAV-файл запроса, указать параметр top-k и отправить форму на сервер."
    ])
    add_body_paragraph(doc, "Такой сценарий важен и в методическом смысле: он отделяет одноразовую подготовку базы от многократных пользовательских запросов. Благодаря этому скорость ответа сервиса определяется только построением query embedding и вычислением скалярных произведений, а не повторной обработкой всей коллекции. В учебной работе это наглядно демонстрирует преимущество индексированного поиска над наивным полным пересчетом.")

    add_heading(doc, "10. Экспериментальная проверка")
    add_body_paragraph(doc, "Эксперименты проводились в двух сценариях. В первом сценарии запросом был один из треков, уже входящих в базу. Это позволяет проверить, находит ли система сам трек на первом месте и насколько близкими считает соседние композиции. Во втором сценарии в качестве запроса использовалась реальная диктофонная запись, что моделирует более сложную и практическую ситуацию, когда запросный аудиосигнал отличается от базы по качеству записи, шумам и общему акустическому домену.")

    add_experiment_table(
        doc,
        "Таблица 3 — Результаты поиска для запроса dataset/rock/rock.00000.wav",
        [
            ("rock.00000", "1.0000", "0.9923"),
            ("rock.00002", "0.4438", "0.8883"),
            ("disco.00000", "0.4189", "0.8753"),
            ("rock.00004", "0.3932", "0.8606"),
            ("hiphop.00000", "0.3556", "0.8364"),
        ],
    )
    add_body_paragraph(doc, "Как видно из таблицы 1, при запросе по самому треку система корректно помещает его на первое место. Это означает, что индекс и поиск работают правильно на уровне инженерного пайплайна. При этом уже вторые и третьи позиции могут относиться не только к жанру rock. Такой результат показывает, что текущий embedding скорее описывает общий акустический профиль, чем строгое жанровое или композиционное соответствие.")

    add_experiment_table(
        doc,
        "Таблица 4 — Результаты поиска для сжатого WAV-файла после компрессии",
        [
            ("jazz.00009", "0.8142", "0.8840"),
            ("reggae.00001", "0.7846", "0.8765"),
            ("classical.00009", "0.7733", "0.8735"),
            ("classical.00006", "0.6653", "0.8417"),
            ("blues.00003", "0.6150", "0.8248"),
        ],
    )
    add_body_paragraph(doc, "Во втором сценарии для проверки устойчивости сервиса был использован не исходный файл из базы, а его сжатая версия. Компрессия WAV-файла выполнялась с помощью онлайн-сервиса FreeConvert WAV Compressor, после чего полученный файл загружался в веб-интерфейс системы как отдельный запрос. Такой эксперимент интересен тем, что позволяет проверить, насколько сильно умеренное изменение аудиосигнала при компрессии влияет на итоговое положение правильного трека в выдаче.")
    add_body_paragraph(doc, "Результат эксперимента показал, что после компрессии запросный файл по-прежнему наиболее близок к исходному треку jazz.00009, который занял первое место в выдаче с similarity 0.8142 и confidence 0.8840. Это говорит о том, что предложенное представление признаков сохраняет достаточную устойчивость к умеренным искажениям сигнала и не разрушается при повторном кодировании аудио. Остальные позиции заняли треки с частично похожей спектральной структурой, однако правильная композиция все равно осталась лидером.")
    add_body_paragraph(doc, "Таким образом, проведенный эксперимент подтверждает работоспособность системы не только на полностью совпадающем файле из базы, но и на преобразованной версии сигнала. Для учебной работы это важный результат, поскольку он демонстрирует практическую устойчивость retrieval-подхода к компрессии и умеренным изменениям аудиоданных.")

    add_heading(doc, "11. Анализ ошибок и интерпретация ограничений")
    add_body_paragraph(doc, "Основная ошибка, выявленная при отладке проекта, заключалась в слишком высокой и плохо интерпретируемой косинусной близости практически для всех объектов базы. Изначально embedding формировался из усредненных статистик, большинство компонентов которых имело близкий положительный масштаб. Это приводило к сжатию пространства признаков и делало similarity почти бесполезной для содержательного сравнения. После анализа распределения оценок была внедрена стандартизация признаков по базе и новая схема confidence, основанная на z-оценке кандидата относительно всех ответов на запрос.")
    add_body_paragraph(doc, "Второй класс ошибок связан с чувствительностью retrieval-системы к способу подготовки аудиофайла. Даже если трек после компрессии или перекодирования визуально кажется тем же самым объектом, его спектральная структура и распределение энергии по полосам меняются. Поэтому важно, чтобы pipeline извлечения признаков был устойчив к таким преобразованиям. В ходе доработки baseline-модели удалось добиться того, что сжатая версия тестового файла все равно корректно возвращала исходный трек на первом месте.")
    add_body_paragraph(doc, "Третий источник ошибок — потеря временной структуры. Текущая модель агрегирует статистики по всей записи и почти не учитывает, как именно развиваются во времени ритм, акценты и музыкальные события. Два сигнала могут иметь похожее распределение энергии по частотам, но принадлежать разным композициям. Следовательно, для более точного matching необходимо переходить к сегментному сравнению, sequence-моделям или специализированным audio fingerprinting-архитектурам.")

    add_heading(doc, "12. Ограничения разработанного решения и пути улучшения")
    add_body_paragraph(doc, "Несмотря на успешную реализацию полного сервиса, текущий проект имеет ряд ограничений. Во-первых, поддерживаются только PCM WAV-файлы. Во-вторых, embedding формируется из агрегированных статистик и потому плохо учитывает длинные временные зависимости, ритмический рисунок и локальные повторяющиеся паттерны. В-третьих, обученные на музыкальных данных признаки в baseline-варианте все равно недостаточно устойчивы к внешним шумам, голосовым вставкам и диктофонным записям.")
    add_body_paragraph(doc, "Наиболее естественные направления развития следующие:")
    add_bullet_list(doc, [
        "поддержка MP3 и других популярных форматов через ffmpeg или librosa;",
        "переход от handcrafted-признаков к pretrained audio embeddings, например PANNs, CLAP или другим self-supervised моделям;",
        "использование сегментного поиска по нескольким окнам внутри трека с голосованием результатов;",
        "замена точного поиска по NumPy на FAISS или HNSW для больших баз;",
        "автоматическая оценка качества retrieval по метрикам Recall@K, MRR и precision."
    ])
    add_body_paragraph(doc, "С инженерной точки зрения полезным будет также развитие интерфейса: добавление drag-and-drop загрузки, просмотра истории запросов, жанровой статистики и отдельной страницы администрирования индекса. Однако даже в текущем виде сервис уже демонстрирует ключевые этапы production-like решения: оффлайн-подготовку, онлайн-поиск, веб-интерфейс, модульную архитектуру и интерпретируемый анализ ошибок.")

    add_heading(doc, "13. Вывод")
    add_body_paragraph(doc, "В ходе работы был разработан сервис поиска похожих музыкальных композиций на языке Python. Реализация охватывает полный цикл обработки аудиоданных: от подготовки датасета GTZAN и извлечения признаков до построения индекса, веб-сервиса и пользовательского интерфейса с загрузкой файла. Базовый алгоритм использует спектральные признаки, мел-энергии, MFCC, стандартизацию и косинусную близость. Дополнительно была выполнена калибровка confidence, позволившая сделать оценки поиска более адекватными и интерпретируемыми.")
    add_body_paragraph(doc, "Экспериментальная часть показала, что система корректно работает на уровне пайплайна, уверенно находит сам трек в базе и сохраняет работоспособность даже после компрессии запросного WAV-файла. Таким образом, поставленная учебная задача решена: создан и протестирован рабочий сервис audio similarity search, а также проведен содержательный анализ его достоинств, ограничений и направлений дальнейшего развития.")

    add_heading(doc, "Список литературы")
    add_number_list(doc, [
        "Tzanetakis G., Cook P. Musical Genre Classification of Audio Signals // IEEE Transactions on Speech and Audio Processing. 2002. Vol. 10, No. 5. P. 293-302.",
        "Davis S., Mermelstein P. Comparison of Parametric Representations for Monosyllabic Word Recognition in Continuously Spoken Sentences // IEEE Transactions on Acoustics, Speech, and Signal Processing. 1980. Vol. 28, No. 4. P. 357-366.",
        "Kong Q., Cao Y., Iqbal T., Wang Y., Wang W., Plumbley M. PANNs: Large-Scale Pretrained Audio Neural Networks for Audio Pattern Recognition // IEEE/ACM Transactions on Audio, Speech, and Language Processing. 2020. Vol. 28. P. 2880-2894.",
        "Johnson J., Douze M., Jegou H. Billion-scale similarity search with GPUs // IEEE Transactions on Big Data. 2019. Vol. 7, No. 3. P. 535-547.",
        "GTZAN Genre Collection [Электронный ресурс]. Режим доступа: открытый набор данных для задач музыкальной классификации и retrieval.",
        "Документация проекта и исходный код сервиса поиска похожих музыкальных композиций, разработанный в рамках лабораторной работы."
    ])

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_report()
    print(output)
